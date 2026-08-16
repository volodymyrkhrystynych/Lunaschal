"""Reading an existing resume into the profile.

The model is stubbed everywhere. What is under test is the bound: the model
returns *line numbers* for accomplishments and never their text, so an imported
bullet is byte-identical to what the document said. That matters more here than
in tailoring — a tailored bullet is reviewed once and sent, while an imported
one becomes `profile_bullets`, which every future resume is generated from and
which the anti-fabrication guarantee treats as fact.
"""
from io import BytesIO

import pytest

from backend.db.connection import get_db
from backend.jobs import resume_import


RESUME_TEXT = """Ada Lovelace
ada@example.com · +1 416 555 0100 · Toronto, ON

EXPERIENCE

Acme Inc — Senior Engineer — 2021–Present
• Built the billing service in Python and Postgres.
• Cut page load time in half.

Initech — Engineer — 2019–2021
• Migrated the monolith to Kubernetes.

SKILLS
Python, Postgres, Kubernetes

EDUCATION
University of Toronto — BSc Computer Science — 2015–2019
"""


def lines_of(text: str = RESUME_TEXT):
    return resume_import.extract_lines(text=text)


def docx_bytes(paragraphs: list[tuple[str, str]]) -> bytes:
    """A real .docx, built in-process. `paragraphs` is [(style, text)]."""
    from docx import Document

    document = Document()
    for style, text in paragraphs:
        document.add_paragraph(text, style=style) if style else document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------
# Extraction
# --------------------------------------------------------------------------

def test_pasted_text_becomes_numbered_lines():
    lines = lines_of()
    assert lines[0].text == 'Ada Lovelace'
    assert [line.index for line in lines] == list(range(len(lines)))
    assert all(line.text.strip() == line.text for line in lines)


def test_blank_lines_are_dropped_so_indexes_stay_dense():
    lines = resume_import.extract_lines(text='One\n\n\n   \nTwo')
    assert [line.text for line in lines] == ['One', 'Two']


def test_a_character_bulleted_line_is_recognised():
    """Plenty of resumes fake a list with '•' instead of using Word's style."""
    lines = resume_import.extract_lines(text='• Built the billing service.')
    assert lines[0].looks_like_bullet is True


def test_a_plain_line_is_not_a_bullet():
    assert resume_import.extract_lines(text='EXPERIENCE')[0].looks_like_bullet is False


def test_a_real_docx_list_paragraph_is_recognised():
    """mammoth maps Word's list style onto <li>, which is the one piece of
    structure worth having — it separates accomplishments from headings
    without guessing."""
    data = docx_bytes([
        ('', 'Acme Inc — Senior Engineer'),
        ('List Bullet', 'Built the billing service in Python.'),
        ('List Bullet', 'Cut page load time in half.'),
    ])
    lines = resume_import.extract_lines(data=data, filename='resume.docx')

    texts = [line.text for line in lines]
    assert 'Built the billing service in Python.' in texts
    listed = [line for line in lines if line.is_list]
    assert len(listed) == 2


def test_a_docx_heading_is_marked_as_one():
    data = docx_bytes([('Heading 1', 'EXPERIENCE'), ('', 'Acme Inc')])
    lines = resume_import.extract_lines(data=data, filename='resume.docx')
    assert lines[0].is_heading is True
    assert lines[1].is_heading is False


def test_a_non_docx_upload_is_refused():
    with pytest.raises(ValueError):
        resume_import.extract_lines(data=b'%PDF-1.4', filename='resume.pdf')


def test_a_corrupt_docx_is_refused_with_a_message():
    with pytest.raises(ValueError):
        resume_import.extract_lines(data=b'not a zip file', filename='resume.docx')


def test_line_count_is_capped(monkeypatch):
    monkeypatch.setattr(resume_import, 'MAX_LINES', 5)
    assert len(resume_import.extract_lines(text='\n'.join(str(i) for i in range(50)))) == 5


# --------------------------------------------------------------------------
# The bound
# --------------------------------------------------------------------------

def test_the_schema_bounds_indexes_to_the_real_lines():
    schema = resume_import.build_schema(12)
    index = schema['properties']['roles']['items']['properties']['bulletIndexes']['items']
    assert index == {'type': 'integer', 'minimum': 0, 'maximum': 11}


def test_the_schema_has_no_field_for_bullet_text():
    """The strongest form of the guarantee: prose is not merely rejected, it is
    unrepresentable."""
    role = resume_import.build_schema(5)['properties']['roles']['items']['properties']
    assert set(role) == {
        'company', 'title', 'location', 'startLabel', 'endLabel', 'bulletIndexes'
    }


def test_an_empty_document_still_produces_a_valid_schema():
    index = (resume_import.build_schema(0)['properties']['roles']['items']
             ['properties']['bulletIndexes']['items'])
    assert index['maximum'] == 0


def test_bullet_text_comes_from_the_document_not_the_model():
    """The central guarantee. Even handed altered prose alongside the index,
    the stored text is what the resume actually said."""
    lines = lines_of()
    target = next(l for l in lines if 'billing service' in l.text)

    result = resume_import.clamp({
        'roles': [{
            'company': 'Acme Inc', 'title': 'Senior Engineer',
            'bulletIndexes': [target.index],
            # A model that tried to smuggle prose through anyway:
            'bullets': ['Architected a world-class billing platform.'],
        }],
    }, lines)

    text = result['roles'][0]['bullets'][0]['text']
    assert text == 'Built the billing service in Python and Postgres.'
    assert 'world-class' not in text


def test_the_bullet_marker_is_stripped_but_nothing_else_is():
    lines = resume_import.extract_lines(text='• Built the billing service.')
    result = resume_import.clamp(
        {'roles': [{'company': 'Acme', 'title': 'Eng', 'bulletIndexes': [0]}]}, lines
    )
    assert result['roles'][0]['bullets'][0]['text'] == 'Built the billing service.'


def test_an_out_of_range_index_is_dropped():
    lines = lines_of()
    result = resume_import.clamp({
        'roles': [{'company': 'Acme', 'title': 'Eng',
                   'bulletIndexes': [0, 9999, -3, 'nonsense']}],
    }, lines)
    assert [b['index'] for b in result['roles'][0]['bullets']] == [0]


def test_one_line_cannot_land_under_two_roles():
    """A duplicated accomplishment reads as padding on the rendered resume."""
    lines = lines_of()
    result = resume_import.clamp({
        'roles': [
            {'company': 'Acme', 'title': 'Eng', 'bulletIndexes': [5]},
            {'company': 'Initech', 'title': 'Eng', 'bulletIndexes': [5]},
        ],
    }, lines)
    assert len(result['roles'][0]['bullets']) == 1
    assert result['roles'][1]['bullets'] == []


def test_a_repeated_index_within_one_role_is_deduped():
    lines = lines_of()
    result = resume_import.clamp(
        {'roles': [{'company': 'Acme', 'title': 'Eng', 'bulletIndexes': [5, 5, 5]}]},
        lines,
    )
    assert len(result['roles'][0]['bullets']) == 1


def test_a_role_with_neither_company_nor_title_is_dropped():
    """Usually the model trying to turn a section heading into a job."""
    result = resume_import.clamp({
        'roles': [
            {'company': '', 'title': '', 'bulletIndexes': [0]},
            {'company': 'Acme', 'title': '', 'bulletIndexes': []},
        ],
    }, lines_of())
    assert [r['company'] for r in result['roles']] == ['Acme']


def test_skills_are_trimmed_and_deduped_case_insensitively():
    result = resume_import.clamp(
        {'roles': [], 'skills': ['Python', ' python ', 'Postgres', '']}, lines_of()
    )
    assert result['skills'] == ['Python', 'Postgres']


def test_education_without_an_institution_or_credential_is_dropped():
    result = resume_import.clamp({
        'roles': [],
        'education': [{'institution': '', 'credential': ''},
                      {'institution': 'University of Toronto', 'credential': 'BSc'}],
    }, lines_of())
    assert len(result['education']) == 1


def test_unused_lines_are_reported_so_nothing_vanishes_silently():
    """The review screen can show what the parser did not place, which is how
    a dropped accomplishment becomes visible instead of just missing."""
    lines = lines_of()
    result = resume_import.clamp(
        {'roles': [{'company': 'Acme', 'title': 'Eng', 'bulletIndexes': [5]}]}, lines
    )
    assert len(result['unusedLines']) == len(lines) - 1
    assert all('index' in u and 'text' in u for u in result['unusedLines'])


def test_import_returns_none_when_the_model_is_unavailable(monkeypatch):
    """None, not an empty profile — the two must never look alike, or the user
    retypes a resume that was actually read fine."""
    monkeypatch.setattr(resume_import, 'is_ai_configured', lambda: False)
    assert resume_import.import_resume(lines_of()) is None


def test_import_returns_none_when_the_model_raises(monkeypatch):
    monkeypatch.setattr(resume_import, 'is_ai_configured', lambda: True)
    monkeypatch.setattr(resume_import, 'chat_json',
                        lambda *a, **k: (_ for _ in ()).throw(RuntimeError('down')))
    assert resume_import.import_resume(lines_of()) is None


def test_import_of_an_empty_document_is_none():
    assert resume_import.import_resume([]) is None


# --------------------------------------------------------------------------
# Through HTTP
# --------------------------------------------------------------------------

@pytest.fixture
def stub_model(monkeypatch):
    """Returns the shape a well-behaved model would, over RESUME_TEXT."""
    def fake(prompt, system=None, schema=None, max_tokens=None, **kwargs):
        return {
            'contact': {'fullName': 'Ada Lovelace', 'email': 'ada@example.com',
                        'phone': '+1 416 555 0100', 'location': 'Toronto, ON'},
            'roles': [
                # Lines 4/5 are Acme's bullets and 7 is Initech's — see the
                # numbering RESUME_TEXT produces once blanks are dropped.
                {'company': 'Acme Inc', 'title': 'Senior Engineer',
                 'startLabel': '2021', 'endLabel': 'Present',
                 'bulletIndexes': [4, 5]},
                {'company': 'Initech', 'title': 'Engineer',
                 'startLabel': '2019', 'endLabel': '2021', 'bulletIndexes': [7]},
            ],
            'skills': ['Python', 'Postgres', 'Kubernetes'],
            'education': [{'institution': 'University of Toronto',
                           'credential': 'BSc', 'field': 'Computer Science'}],
        }

    monkeypatch.setattr(resume_import, 'chat_json', fake)
    monkeypatch.setattr(resume_import, 'is_ai_configured', lambda: True)


def test_preview_writes_nothing(client, stub_model):
    """Preview-then-commit is the point: this feeds the one table the
    anti-fabrication guarantee treats as fact."""
    body = client.post('/api/jobs/profile/import',
                       json={'text': RESUME_TEXT}).get_json()

    assert len(body['roles']) == 2
    assert body['roles'][0]['company'] == 'Acme Inc'
    assert get_db().execute('SELECT COUNT(*) AS c FROM profile_roles').fetchone()['c'] == 0


def test_preview_accepts_a_docx_upload(client, stub_model):
    data = docx_bytes([
        ('', 'Acme Inc — Senior Engineer'),
        ('List Bullet', 'Built the billing service in Python.'),
    ])
    response = client.post(
        '/api/jobs/profile/import',
        data={'file': (BytesIO(data), 'resume.docx')},
        content_type='multipart/form-data',
    )
    assert response.status_code == 200


def test_preview_refuses_a_pdf(client, stub_model):
    response = client.post(
        '/api/jobs/profile/import',
        data={'file': (BytesIO(b'%PDF-1.4'), 'resume.pdf')},
        content_type='multipart/form-data',
    )
    assert response.status_code == 400
    assert 'paste' in response.get_json()['error'].lower()


def test_preview_needs_something_to_read(client, stub_model):
    assert client.post('/api/jobs/profile/import', json={}).status_code == 400


def test_preview_is_503_when_the_model_is_down(client, monkeypatch):
    monkeypatch.setattr(resume_import, 'is_ai_configured', lambda: False)
    response = client.post('/api/jobs/profile/import', json={'text': RESUME_TEXT})

    assert response.status_code == 503
    assert get_db().execute('SELECT COUNT(*) AS c FROM profile_roles').fetchone()['c'] == 0


def test_preview_releases_its_priority_mark(client, monkeypatch):
    from backend.ai import priority

    monkeypatch.setattr(resume_import, 'is_ai_configured', lambda: True)
    monkeypatch.setattr(resume_import, 'chat_json',
                        lambda *a, **k: (_ for _ in ()).throw(RuntimeError('boom')))
    client.post('/api/jobs/profile/import', json={'text': RESUME_TEXT})

    assert priority.active() is False


def test_commit_writes_the_reviewed_structure(client, stub_model):
    preview = client.post('/api/jobs/profile/import',
                          json={'text': RESUME_TEXT}).get_json()

    result = client.post('/api/jobs/profile/import/commit', json=preview)

    assert result.status_code == 201
    created = result.get_json()['created']
    assert created == {'roles': 2, 'bullets': 3, 'skills': 3, 'education': 1}

    db = get_db()
    rows = db.execute('SELECT text FROM profile_bullets ORDER BY ord').fetchall()
    assert any('billing service' in r['text'] for r in rows)


def test_commit_sets_ord_in_order(client, stub_model):
    preview = client.post('/api/jobs/profile/import',
                          json={'text': RESUME_TEXT}).get_json()
    client.post('/api/jobs/profile/import/commit', json=preview)

    roles = get_db().execute(
        'SELECT company, ord FROM profile_roles ORDER BY ord'
    ).fetchall()
    assert [r['company'] for r in roles] == ['Acme Inc', 'Initech']
    assert [r['ord'] for r in roles] == [0, 1]


def test_commit_appends_rather_than_replacing(client, stub_model):
    client.post('/api/jobs/profile/roles', json={'company': 'Existing', 'title': 'Dev'})
    preview = client.post('/api/jobs/profile/import',
                          json={'text': RESUME_TEXT}).get_json()

    client.post('/api/jobs/profile/import/commit', json=preview)

    companies = [
        r['company'] for r in
        get_db().execute('SELECT company FROM profile_roles ORDER BY ord').fetchall()
    ]
    assert 'Existing' in companies
    assert 'Acme Inc' in companies


def test_commit_does_not_overwrite_a_contact_field_already_filled(client, stub_model):
    """The user may have corrected a phone number by hand; an import must not
    undo that."""
    client.patch('/api/jobs/profile', json={'phone': '+1 647 555 9999'})
    preview = client.post('/api/jobs/profile/import',
                          json={'text': RESUME_TEXT}).get_json()

    client.post('/api/jobs/profile/import/commit', json=preview)

    profile = client.get('/api/jobs/profile').get_json()['profile']
    assert profile['phone'] == '+1 647 555 9999'
    assert profile['fullName'] == 'Ada Lovelace'   # blank before, so filled


def test_commit_skips_a_skill_already_on_the_profile(client, stub_model):
    client.post('/api/jobs/profile/skills', json={'name': 'python'})
    preview = client.post('/api/jobs/profile/import',
                          json={'text': RESUME_TEXT}).get_json()

    created = client.post('/api/jobs/profile/import/commit',
                          json=preview).get_json()['created']

    assert created['skills'] == 2   # Postgres and Kubernetes; Python was there


def test_commit_makes_no_model_call(client, monkeypatch):
    """It stores exactly what the user approved — nothing is re-derived."""
    def explode(*a, **k):
        raise AssertionError('commit must not call the model')

    monkeypatch.setattr('backend.ai.llm.chat_json', explode)
    response = client.post('/api/jobs/profile/import/commit', json={
        'roles': [{'company': 'Acme', 'title': 'Eng',
                   'bullets': [{'text': 'Did the thing.'}]}],
        'skills': ['Python'],
    })
    assert response.status_code == 201


def test_commit_ignores_a_role_the_user_emptied(client):
    response = client.post('/api/jobs/profile/import/commit', json={
        'roles': [{'company': '', 'title': '', 'bullets': [{'text': 'x'}]}],
    })
    assert response.get_json()['created']['roles'] == 0


def test_the_import_route_is_not_shadowed_by_the_child_factory(client, stub_model):
    """`/profile/import` and `/profile/<kind>` are both POST; the static rule
    has to win or an import would try to create a child row of kind 'import'."""
    body = client.post('/api/jobs/profile/import',
                       json={'text': RESUME_TEXT}).get_json()
    assert 'roles' in body and 'id' not in body


# --------------------------------------------------------------------------
# What the whole thing is for
# --------------------------------------------------------------------------

def test_importing_a_resume_makes_the_feed_sortable(client, stub_model):
    """Before a profile exists every posting scores NULL and the feed cannot
    sort. This is the moment that stops being true."""
    job_id = client.post('/api/jobs', json={
        'title': 'Backend Engineer', 'company': 'Someco',
        'description': 'We need Python and Postgres experience.',
    }).get_json()['id']
    assert client.get('/api/jobs/feed').get_json()[0]['matchScore'] is None

    preview = client.post('/api/jobs/profile/import',
                          json={'text': RESUME_TEXT}).get_json()
    client.post('/api/jobs/profile/import/commit', json=preview)
    client.post('/api/jobs/rescore')

    scored = client.get('/api/jobs/feed').get_json()[0]
    assert scored['id'] == job_id
    assert scored['matchScore'] is not None and scored['matchScore'] > 0
