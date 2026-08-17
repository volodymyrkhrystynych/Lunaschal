"""Rendering a tailored resume to HTML, PDF and DOCX.

One HTML template is the single source of layout: the app previews that exact
string and WeasyPrint prints it, so what you approve on screen is what the
employer opens. The DOCX is built separately because a Word file is a different
object model, not a rendering of HTML — but it is built from the same content
dict, so the two cannot disagree about *content*, only about typography.

The layout is deliberately plain: one column, no tables, no text boxes, no
sidebars, standard section headings. Résumé screeners read documents by
flattening them, and a two-column layout flattens into interleaved nonsense.
The prettiest resume that parses wrong scores zero.

Both renderers are imported lazily. A missing WeasyPrint costs you the PDF and
nothing else — the HTML preview, the DOCX and the whole rest of the feature
keep working, and `is_pdf_available()` lets the UI say so plainly.
"""
import html as html_mod
import logging
import re
import unicodedata
from pathlib import Path

logger = logging.getLogger(__name__)

# Anything that would either break out of a filename or break the header it
# travels in. Path separators are the obvious half; the control characters
# matter because the name lands in a Content-Disposition header, and a bare
# CR/LF there is header injection. Werkzeug quotes the value, but a sanitizer
# that depends on the framework's escaping is one refactor from being wrong.
_FILENAME_STRIP = re.compile(r'[\\/:*?"<>|\x00-\x1f\x7f]')
_FILENAME_SPACE = re.compile(r'\s+')
_MAX_STEM = 80


def download_filename(full_name: str, ext: str) -> str:
    """The name the employer sees on the file: "Ada Lovelace Resume.pdf".

    Recruiters receive hundreds of these, and a mailbox of `resume.pdf` is
    exactly as useless as it sounds. Falls back to a bare "Resume" when the
    profile has no name, which is better than an empty stem.
    """
    # NFC first, so an accented name composed of combining characters comes out
    # as one character per glyph rather than surviving as a decomposed pair.
    cleaned = _FILENAME_STRIP.sub(' ', unicodedata.normalize('NFC', full_name or ''))
    cleaned = _FILENAME_SPACE.sub(' ', cleaned).strip(' .')[:_MAX_STEM].strip()
    stem = f'{cleaned} Resume' if cleaned else 'Resume'
    return f'{stem}.{ext}'


def is_pdf_available() -> bool:
    try:
        import weasyprint  # noqa: F401
    except Exception:
        return False
    return True


def is_docx_available() -> bool:
    try:
        import docx  # noqa: F401
    except Exception:
        return False
    return True


_CSS = """
:root { color-scheme: light; }
.resume {
  font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
  font-size: 10.5pt;
  line-height: 1.4;
  color: #111;
  background: #fff;
  max-width: 7.5in;
  margin: 0 auto;
}
.resume h1 { font-size: 20pt; margin: 0 0 2pt; letter-spacing: 0.2pt; }
.resume .headline { font-size: 11pt; color: #333; margin: 0 0 4pt; }
.resume .contact { font-size: 9.5pt; color: #333; margin: 0 0 12pt; }
.resume .contact span + span::before { content: " · "; color: #999; }
.resume h2 {
  font-size: 10.5pt; text-transform: uppercase; letter-spacing: 0.8pt;
  border-bottom: 1px solid #999; padding-bottom: 2pt;
  margin: 14pt 0 7pt;
}
.resume .role { margin-bottom: 9pt; }
.resume .role-head {
  display: flex; justify-content: space-between; gap: 12pt;
  align-items: baseline;
}
.resume .role-title { font-weight: 700; }
.resume .role-company { font-weight: 400; }
.resume .role-dates { font-size: 9.5pt; color: #444; white-space: nowrap; }
.resume ul { margin: 3pt 0 0; padding-left: 15pt; }
.resume li { margin-bottom: 2.5pt; }
.resume .summary { margin: 0; }
.resume .skills { margin: 0; }
.resume .skill-group { margin-bottom: 3pt; }
.resume .skill-group .label { font-weight: 700; }
.resume .edu { margin-bottom: 5pt; }
@page { size: Letter; margin: 0.6in 0.7in; }
"""


def _esc(value) -> str:
    return html_mod.escape(str(value or ''), quote=True)


def _dates(role: dict) -> str:
    start = (role.get('startLabel') or '').strip()
    end = (role.get('endLabel') or '').strip()
    if start and end:
        return f'{start} – {end}'
    return start or end


def _selected_by_role(content: dict) -> dict[str, list[dict]]:
    grouped: dict[str, list[dict]] = {}
    for item in content.get('selectedBullets') or []:
        grouped.setdefault(item.get('roleId') or '', []).append(item)
    return grouped


def _ordered_skills(loaded: dict, content: dict) -> list[dict]:
    """Skills with the posting's matched keywords pulled to the front.

    A screener — human or machine — reads the first line of the skills block
    and often no further, so the terms this posting actually asked for belong
    there rather than wherever they happened to be typed.
    """
    matched = {m.lower() for m in (content.get('keywords') or {}).get('matched') or []}
    skills = list(loaded.get('skills') or [])
    return sorted(skills, key=lambda s: (0 if (s.get('name') or '').lower() in matched else 1,))


def render_html(loaded: dict, content: dict, job: dict | None = None) -> str:
    """The resume as a self-contained HTML fragment plus its <style>."""
    profile = loaded.get('profile') or {}
    grouped = _selected_by_role(content)

    contact_bits = [
        profile.get('email'), profile.get('phone'), profile.get('location'),
        *[link.get('url') for link in (profile.get('links') or [])
          if isinstance(link, dict) and link.get('url')],
    ]
    contact = ''.join(f'<span>{_esc(b)}</span>' for b in contact_bits if b)

    parts = [f'<style>{_CSS}</style>', '<div class="resume">']
    parts.append(f'<h1>{_esc(profile.get("fullName"))}</h1>')
    if profile.get('headline'):
        parts.append(f'<p class="headline">{_esc(profile["headline"])}</p>')
    if contact:
        parts.append(f'<p class="contact">{contact}</p>')

    summary = content.get('summary') or profile.get('summary')
    if summary:
        parts.append('<h2>Summary</h2>')
        parts.append(f'<p class="summary">{_esc(summary)}</p>')

    roles = loaded.get('roles') or []
    if roles:
        parts.append('<h2>Experience</h2>')
        for role in roles:
            bullets = grouped.get(role.get('id'), [])
            parts.append('<div class="role">')
            parts.append('<div class="role-head"><div>')
            parts.append(f'<span class="role-title">{_esc(role.get("title"))}</span>')
            if role.get('company'):
                parts.append(f'<span class="role-company"> — {_esc(role["company"])}</span>')
            parts.append('</div>')
            dates = _dates(role)
            if dates:
                parts.append(f'<div class="role-dates">{_esc(dates)}</div>')
            parts.append('</div>')
            if bullets:
                parts.append('<ul>')
                parts.extend(f'<li>{_esc(b.get("text"))}</li>' for b in bullets)
                parts.append('</ul>')
            parts.append('</div>')

    skills = _ordered_skills(loaded, content)
    if skills:
        parts.append('<h2>Skills</h2>')
        by_category: dict[str, list[str]] = {}
        for skill in skills:
            by_category.setdefault(skill.get('category') or '', []).append(skill.get('name') or '')
        parts.append('<div class="skills">')
        for category, names in by_category.items():
            label = f'<span class="label">{_esc(category)}: </span>' if category else ''
            parts.append(f'<div class="skill-group">{label}{_esc(", ".join(n for n in names if n))}</div>')
        parts.append('</div>')

    education = loaded.get('education') or []
    if education:
        parts.append('<h2>Education</h2>')
        for edu in education:
            line = ' — '.join(x for x in [edu.get('credential'), edu.get('field')] if x)
            parts.append('<div class="edu">')
            parts.append(f'<div><strong>{_esc(edu.get("institution"))}</strong></div>')
            tail = ' '.join(x for x in [line, _dates(edu)] if x)
            if tail:
                parts.append(f'<div>{_esc(tail)}</div>')
            if edu.get('notes'):
                parts.append(f'<div>{_esc(edu["notes"])}</div>')
            parts.append('</div>')

    parts.append('</div>')
    return '\n'.join(parts)


def render_pdf(html: str, path: Path) -> bool:
    """Write the PDF. False when WeasyPrint isn't installed or the write fails."""
    try:
        from weasyprint import HTML
    except Exception as e:
        logger.info('PDF rendering unavailable: %s', e)
        return False
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        HTML(string=f'<!doctype html><meta charset="utf-8">{html}').write_pdf(str(path))
        return True
    except Exception as e:
        logger.warning('PDF rendering failed: %s', e)
        return False


def render_docx(loaded: dict, content: dict, path: Path) -> bool:
    """Write the .docx. False when python-docx isn't installed or it fails."""
    try:
        import docx
        from docx.shared import Pt
    except Exception as e:
        logger.info('DOCX rendering unavailable: %s', e)
        return False

    try:
        profile = loaded.get('profile') or {}
        grouped = _selected_by_role(content)
        document = docx.Document()

        normal = document.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10.5)

        document.add_heading(profile.get('fullName') or '', level=0)
        if profile.get('headline'):
            document.add_paragraph(profile['headline'])

        contact_bits = [
            profile.get('email'), profile.get('phone'), profile.get('location'),
            *[link.get('url') for link in (profile.get('links') or [])
              if isinstance(link, dict) and link.get('url')],
        ]
        contact = ' · '.join(b for b in contact_bits if b)
        if contact:
            document.add_paragraph(contact)

        summary = content.get('summary') or profile.get('summary')
        if summary:
            document.add_heading('Summary', level=1)
            document.add_paragraph(summary)

        roles = loaded.get('roles') or []
        if roles:
            document.add_heading('Experience', level=1)
            for role in roles:
                head = ' — '.join(x for x in [role.get('title'), role.get('company')] if x)
                dates = _dates(role)
                paragraph = document.add_paragraph()
                paragraph.add_run(head).bold = True
                if dates:
                    paragraph.add_run(f'   {dates}')
                for bullet in grouped.get(role.get('id'), []):
                    document.add_paragraph(bullet.get('text') or '', style='List Bullet')

        skills = _ordered_skills(loaded, content)
        if skills:
            document.add_heading('Skills', level=1)
            by_category: dict[str, list[str]] = {}
            for skill in skills:
                by_category.setdefault(skill.get('category') or '', []).append(skill.get('name') or '')
            for category, names in by_category.items():
                joined = ', '.join(n for n in names if n)
                paragraph = document.add_paragraph()
                if category:
                    paragraph.add_run(f'{category}: ').bold = True
                paragraph.add_run(joined)

        education = loaded.get('education') or []
        if education:
            document.add_heading('Education', level=1)
            for edu in education:
                paragraph = document.add_paragraph()
                paragraph.add_run(edu.get('institution') or '').bold = True
                tail = ' '.join(x for x in [
                    ' — '.join(y for y in [edu.get('credential'), edu.get('field')] if y),
                    _dates(edu),
                ] if x)
                if tail:
                    paragraph.add_run(f'\n{tail}')
                if edu.get('notes'):
                    paragraph.add_run(f'\n{edu["notes"]}')

        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(path))
        return True
    except Exception as e:
        logger.warning('DOCX rendering failed: %s', e)
        return False
